from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/Silver_Voice_STT_Forced_Alignment_진행기록.docx")


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "0B2545",
    "light_blue": "E8EEF5",
    "light_gray": "F2F4F7",
    "callout": "F4F6F9",
    "border": "B7C9DB",
    "white": "FFFFFF",
}


def set_run_font(run, *, name="Calibri", east_asia="Malgun Gothic", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, *, name="Calibri", east_asia="Malgun Gothic", size=None, bold=None, color=None):
    font = style.font
    font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold=False, color="000000", size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="B7C9DB", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("Silver Voice STT 학습 데이터 정렬 및 Forced Alignment 적용 기록")
    set_run_font(run, size=20, bold=True, color=COLORS["ink"])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("작성일: 2026-07-01 | 목적: 60세 이상 고령자 발화 STT fine-tuning 모델 학습 데이터 개선")
    set_run_font(run, size=10.5, color="555555")


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=11, color="000000")
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_code_block(doc: Document, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="DADCE0", size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F9FB")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, name="Consolas", east_asia="Malgun Gothic", size=9.2, color="1F2937")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(1.7, 4.8)):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    set_table_widths(table, widths)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "항목", bold=True, color=COLORS["ink"])
    set_cell_text(hdr[1], "내용", bold=True, color=COLORS["ink"])
    shade_cell(hdr[0], COLORS["light_gray"])
    shade_cell(hdr[1], COLORS["light_gray"])
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True, color=COLORS["dark_blue"])
        set_cell_text(cells[1], value, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, COLORS["light_blue"])
        set_cell_text(cell, header, bold=True, color=COLORS["ink"], size=9.7)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="B7C9DB", size="6")
    cell = table.cell(0, 0)
    shade_cell(cell, COLORS["callout"])
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=COLORS["dark_blue"])
    p.add_run("\n")
    r = p.add_run(body)
    set_run_font(r, size=10.2, color="000000")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    set_style_font(styles["Heading 1"], size=16, bold=True, color=COLORS["blue"])
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)

    set_style_font(styles["Heading 2"], size=13, bold=True, color=COLORS["blue"])
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)

    set_style_font(styles["Heading 3"], size=12, bold=True, color=COLORS["dark_blue"])
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    for style_name in ("List Bullet", "List Number"):
        set_style_font(styles[style_name], size=10.5)
        styles[style_name].paragraph_format.space_after = Pt(4)
        styles[style_name].paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Silver Voice STT 학습 기록")
    set_run_font(run, size=9, color="666666")


def build_document():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "핵심 결론",
        "현재 문제의 본질은 모델 크기보다 학습 데이터 정렬이다. 평가용으로는 WAV와 라벨을 파일명 또는 sample ID로 1:1 매칭해 CER/WER을 계산하는 방식이 타당하지만, fine-tuning 학습용으로는 answer 음성 구간과 answer 텍스트가 짧고 정확하게 맞는 형태가 필요하다.",
    )

    add_heading(doc, "1. 프로젝트 목표", 1)
    add_body(
        doc,
        "Silver Voice STT의 최종 목표는 프론트엔드나 백엔드 자체가 아니라, 60세 이상 고령자의 발화를 더 정확하게 텍스트로 변환할 수 있는 한국어 STT fine-tuning 모델을 만드는 것이다. 프론트엔드와 백엔드는 학습된 모델의 성능을 확인하기 위한 테스트 환경으로 본다.",
    )
    add_kv_table(
        doc,
        [
            ("입력", "60세 이상 고령자의 실제 발화 음성"),
            ("출력", "정확한 한국어 STT 텍스트"),
            ("기반 모델", "Whisper 계열, 우선 openai/whisper-large-v3"),
            ("학습 방식", "LoRA fine-tuning"),
            ("목표 산출물", "base Whisper + 고령자 음성 특화 LoRA adapter"),
        ],
    )

    add_heading(doc, "2. 현재까지의 주요 실험 결과", 1)
    add_matrix_table(
        doc,
        ["구분", "데이터/환경", "결과", "해석"],
        [
            [
                "CPU 초기 학습",
                "노트북 CPU, 약 2GB 수준 실험 데이터",
                "WER 약 123%, CER 약 94%",
                "학습 파이프라인 확인 목적. 성능 확보 단계는 아니었음.",
            ],
            [
                "NAS whisper-small",
                "NAS RTX 5090, train 약 10,637개",
                "test WER 약 197%, CER 약 151%",
                "GPU 학습과 저장은 성공했지만 데이터 정렬 품질이 낮았음.",
            ],
            [
                "large-v3 baseline",
                "원본 openai/whisper-large-v3, 100개 샘플",
                "WER 약 81.2%, CER 약 64.9%",
                "small보다 출발점이 낫지만 고령자 데이터 특화 필요.",
            ],
            [
                "라벨 검증",
                "전체 라벨 2,923개 중 WAV 매칭 799개",
                "검증 통과 567개, 평균 CER 약 44.99%",
                "학습 전 라벨-음성 품질 필터링 단계로 활용.",
            ],
        ],
        [1.25, 1.8, 1.45, 2.0],
    )

    add_heading(doc, "3. 현재 문제 사항", 1)
    add_body(doc, "AI-Hub 원천 WAV에는 질문자와 답변자의 음성이 같이 들어 있지만, 우리의 목표는 고령자 답변자의 발화를 정확하게 인식하는 모델을 만드는 것이다.")
    add_bullet(doc, "WAV 구조: 질문자 음성 + 60세 이상 답변자 음성")
    add_bullet(doc, "라벨 구조: qa[].question + qa[].answer")
    add_bullet(doc, "부족한 정보: 각 question/answer가 실제 WAV의 몇 초부터 몇 초까지인지에 대한 정밀 timestamp")
    add_bullet(doc, "학습 장애: full 라벨 텍스트가 길어 Whisper decoder의 약 448개 target token 제한을 자주 초과")
    add_bullet(doc, "결과: 검증 통과 manifest 567개 중 실제 학습에는 157개만 사용됨")

    add_kv_table(
        doc,
        [
            ("train", "453개 중 318개가 440토큰 초과"),
            ("valid", "57개 중 44개가 440토큰 초과"),
            ("test", "57개 중 47개가 440토큰 초과"),
            ("결론", "full WAV + full 라벨 방식은 평가에는 가능하지만 fine-tuning 학습 단위로는 부적합"),
        ],
    )

    add_heading(doc, "4. 평가 방식과 학습 방식의 분리", 1)
    add_body(doc, "파일명 또는 sample ID 기준으로 WAV와 라벨을 1:1 매칭하고 STT 결과와 정답 문장을 정규화해 CER/WER을 계산하는 방식은 평가와 라벨 검증에는 적합하다. 하지만 긴 WAV 전체와 긴 라벨 전체를 그대로 학습에 넣으면 토큰 제한과 음성-텍스트 불일치가 생긴다.")
    add_matrix_table(
        doc,
        ["용도", "권장 방식", "비고"],
        [
            ["평가/검증", "WAV 전체 + 라벨 전체를 파일 단위로 비교", "baseline, 라벨 품질 검증, 학습 전후 비교에 적합"],
            ["fine-tuning 학습", "짧은 answer 음성 clip + 해당 answer 텍스트", "토큰 제한과 질문자 음성 혼입 문제를 줄임"],
        ],
        [1.4, 2.8, 2.3],
    )

    add_heading(doc, "5. Forced Alignment 적용 방향", 1)
    add_body(
        doc,
        "Forced alignment는 STT 결과와 라벨을 비교하는 방식이 아니라, 이미 알고 있는 정답 텍스트를 원본 오디오에 맞춰 각 단어 또는 문자 단위의 timestamp를 계산하는 방식이다. 이 방식으로 question/answer turn별 start/end를 추출하고, answer 구간만 학습 clip으로 만들 수 있다.",
    )
    for step in [
        "라벨 JSON의 qa 순서대로 question/answer 텍스트를 하나의 문자열로 연결한다.",
        "원본 WAV 전체와 이미 알고 있는 라벨 텍스트를 forced alignment 도구에 넣는다.",
        "각 단어 또는 문자에 대한 timestamp를 계산한다.",
        "qa turn 경계에 맞춰 question/answer별 start/end를 추출한다.",
        "answer turn만 오디오 clip으로 저장한다.",
        "answer clip + answer text 형태의 학습 manifest를 생성한다.",
    ]:
        add_number(doc, step)

    add_callout(
        doc,
        "기대 효과",
        "질문자 음성을 학습 정답에서 제외하고, 긴 full 라벨을 짧은 answer 단위로 분리할 수 있다. 결과적으로 440토큰 초과 문제를 줄이고, 한 WAV에서 여러 개의 학습 샘플을 확보할 수 있다.",
    )

    add_heading(doc, "6. Bundle에 추가한 파일", 1)
    add_matrix_table(
        doc,
        ["파일", "역할"],
        [
            ["services/training_pipeline/scripts/prepare_forced_aligned_aihub_dataset.py", "WAV + 라벨 텍스트로 timestamp를 만들고 answer-only clip manifest를 생성"],
            ["scripts/prepare-nas-forced-answer-manifest.sh", "NAS에서 forced alignment manifest 생성을 실행하는 wrapper"],
            ["scripts/check-nas-forced-alignment-env.sh", "torchaudio, MMS_FA, CUDA 환경 확인"],
            ["services/training_pipeline/requirements-forced-alignment.txt", "선택 의존성 torchaudio 기록"],
        ],
        [3.5, 3.0],
    )

    add_heading(doc, "7. NAS 실행 절차", 1)
    add_body(doc, "새 ZIP을 NAS에 업로드한 뒤 기존 bundle을 백업하고 새 bundle을 압축 해제한다.")
    add_code_block(
        doc,
        [
            "cd ~/nas_private",
            "OLD=nas_training_bundle_before_forced_align_$(date +%Y%m%d_%H%M%S)",
            "mv nas_training_bundle \"$OLD\"",
            "python3 -m zipfile -e nas_training_bundle.zip nas_training_bundle",
            "mkdir -p nas_training_bundle/reports",
            "cp -r \"$OLD/reports/label_validation\" nas_training_bundle/reports/ 2>/dev/null || true",
            "cd nas_training_bundle",
            "source ~/nas_private/stt-venv/bin/activate",
        ],
    )
    add_body(doc, "forced alignment 환경을 확인한다.")
    add_code_block(doc, ["bash scripts/check-nas-forced-alignment-env.sh"])
    add_body(doc, "torchaudio가 없으면 torch를 교체하지 않도록 --no-deps로 설치한다.")
    add_code_block(
        doc,
        [
            "python3 -m pip install --no-deps torchaudio",
            "bash scripts/check-nas-forced-alignment-env.sh",
        ],
    )
    add_body(doc, "먼저 3개 샘플만 생성해 clip 경계가 자연스러운지 확인한다.")
    add_code_block(
        doc,
        [
            "LIMIT=3 \\",
            "OUTPUT_DIR=\"processed/modern_story_forced_answer_test_v1\" \\",
            "bash scripts/prepare-nas-forced-answer-manifest.sh",
        ],
    )
    add_body(doc, "생성 결과를 확인한다.")
    add_code_block(
        doc,
        [
            "cat processed/modern_story_forced_answer_test_v1/dataset_summary.json",
            "wc -l processed/modern_story_forced_answer_test_v1/train.jsonl",
            "find processed/modern_story_forced_answer_test_v1/clips -name \"*.wav\" | head -10",
        ],
    )

    add_heading(doc, "8. 검수 기준과 다음 단계", 1)
    add_bullet(doc, "LIMIT=3 또는 LIMIT=5로 생성한 answer clip을 직접 들어보고 시작/끝 경계가 자연스러운지 확인한다.")
    add_bullet(doc, "empty_alignment_text 또는 alignment_failed가 많이 나오면 uroman romanization 또는 alignment 전처리 방식을 보완한다.")
    add_bullet(doc, "clip 경계가 안정적이면 검증 통과 567개 전체에 forced alignment를 적용한다.")
    add_bullet(doc, "생성된 answer-only manifest의 token length 분포를 다시 확인한다.")
    add_bullet(doc, "먼저 small 또는 medium으로 빠른 학습 검증을 하고, 성능 개선이 보이면 large-v3 LoRA로 확장한다.")

    add_callout(
        doc,
        "현재 의사결정",
        "모델을 무작정 더 크게 돌리기보다, 학습 데이터의 음성-텍스트 정렬 품질을 먼저 개선한다. 이 과정이 60세 이상 고령자 STT fine-tuning 모델의 성능 개선에 가장 직접적으로 연결된다.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_document()
